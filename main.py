from PyQt6.QtWidgets import QApplication , QWidget,QVBoxLayout,QPushButton,QCheckBox,QLineEdit,QLabel,QMessageBox,QItemDelegate, \
                                                        QTreeWidgetItem,QTreeWidget,QTextEdit,QFrame,QFileDialog,QScrollArea,QMainWindow,QComboBox,QTableWidget,QTableWidgetItem,QRadioButton,QHeaderView,QProgressBar
from PyQt6.QtGui import QIcon,QFont,QIntValidator,QScreen,QFont,QPixmap,QCursor
from PyQt6 import uic
from PyQt6.QtCore import Qt
import sqlite3
import sys
import os
from PIL import Image , ImageOps
import docx
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.oxml.ns import qn as qn2
import webbrowser
import shutil
import ctypes
import pyautogui
import time
from PyPDF2 import PdfMerger
from urllib.parse import quote
import convert_numbers
from docx2pdf import convert
from stuff import suppress_output
xForImpo = 900
yForImpo = 0


title = "موثق البرامج"
icon = "logo.ico"
desktopPath = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False
if is_admin():
    class LoadingPage(QWidget):
        def __init__(self):
            super().__init__()
            uic.loadUi("loadingPage.ui",self)
    class LineEditDelegate(QItemDelegate):
        def createEditor(self, parent, option, index):
            editor = QLineEdit(parent)
            editor.setFrame(False)
            editor.setReadOnly(True)
            editor.setAlignment(Qt.AlignmentFlag.AlignCenter)
            return editor
    class Choices(QWidget):
        def __init__(self):
            super().__init__()
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon(icon))

    con = sqlite3.connect("app.db")
    cr = con.cursor()
    class Main(QMainWindow): #QWidget
        def __init__(self):
            super().__init__()
            self.startWindow = Choices()
            self.setWindowIcon(QIcon("icon.ico"))
            uic.loadUi("desighn.ui",self.startWindow)
            self.startWindow.setWindowIcon(QIcon("icon.ico"))
            self.startWindow.setWindowTitle(title)
            self.startWindow.companyLogo.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
            self.startWindow.savedReports.clicked.connect(self.savedReportsFun)
            self.startWindow.createReports.clicked.connect(self.createReportFun)
            self.startWindow.ControlPanel.clicked.connect(self.controlPanel)
            self.startWindow.savedReports.setStyleSheet("background-color: cyan")
            self.startWindow.createReports.setStyleSheet("background-color: cyan")
            self.startWindow.ControlPanel.setStyleSheet("background-color:rgb(148, 148, 148);")
            self.startWindow.companyLogo.clicked.connect(self.openwebSite)
            self.startWindow.exportDb.clicked.connect(self.exportDb)
            self.startWindow.importDb.clicked.connect(self.importDb)
            self.startWindow.label.setStyleSheet("font-size:30px;")
            self.startWindow.setFixedSize(651,499)
            self.resizeEvent = self.resizedWindow
            self.startWindow.setStyleSheet("font-size:20px;")
            
            self.startWindow.show()

        def openwebSite(self):
            webbrowser.open('https://www.ersal-m.com', new=2)
        def exportDb(self):
            filePath = QFileDialog.getExistingDirectory(self,"Select a Directory",desktopPath)
            if len(filePath)> 0:
                shutil.copy2("app.db",f"{filePath}")
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def importDb(self):
            d = QMessageBox(parent=self.startWindow,text="هل انت متأكد من استبدال قاعدة البيانات؟")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle(title)
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                self.completeImportDb()
        def completeImportDb(self):
            fileDbUser = QFileDialog.getOpenFileName(self.startWindow,"Select a File",desktopPath,filter="Database File (*.db)")
            if len(fileDbUser[0]) > 0:
                try:
                    self.con1 = sqlite3.connect(fileDbUser[0])
                    cr1 = self.con1.cursor()
                    cr1.execute("SELECT useAble FROM confirmationDatabase")
                    if cr1.fetchone()[0] == "canUse":
                        cr1.execute("SELECT * FROM reports")
                        for i in cr1.fetchall():
                            cr.execute("""INSERT INTO reports (reportName , name , Goals , description , executer , executeDate , benefits , countBenefits , pic1 , pic2 , pic3 , pic4 , picLogo , label1Maybe , label2Maybe , manger , co_manger) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""" , (i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8],i[9],i[10],i[11],i[12],i[13],i[14],i[15],i[16],i[17]))                          
                        con.commit()
                        con.close()
                        self.con1.close()
                        d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
                        app.closeAllWindows()
                    else:
                        raise Exception("notUseAble")
                except:
                    self.con1.close()
                    d = QMessageBox(parent=self,text="قاعدة البيانات غير صالحة")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()

        def savedReportsFun(self):
            self.windowSaved = Choices()
            self.windowSaved.setFixedSize(600,500)
            self.windowSaved.setWindowTitle(title)
            self.windowSaved.setWindowIcon(QIcon("icon.ico"))

            self.savedReports = QTableWidget(self.windowSaved)
            self.savedReports.setGeometry(15,5,570,450)

            self.savedReports.setColumnCount(4)
            self.savedReports.setColumnHidden(0,True)
            self.savedReports.setColumnWidth(0,70)
            self.savedReports.setColumnWidth(1,70)
            self.savedReports.setColumnWidth(2,70)
            self.savedReports.setColumnWidth(3,410)
            self.savedReports.setHorizontalHeaderLabels(["","","","اسم التقرير"])
    
            cr.execute("SELECT id,reportName FROM reports")
            for n,i in enumerate(cr.fetchall()):
                self.savedReports.insertRow(self.savedReports.rowCount())
                icon = QPixmap("trashicon.png")
                button = QPushButton()
                button.setStyleSheet(f"Qproperty-icon:url(trashicon.png);qproperty-iconSize:30px 30px;background-color:rgb(253, 253, 253)")
                button.clicked.connect(lambda x,row=n:self.deleteReport(row))
                self.savedReports.setIndexWidget(self.savedReports.model().index(n,1),button)
                self.savedReports.setItem(n,3,QTableWidgetItem(i[1]))
                self.savedReports.setItem(n,0,QTableWidgetItem(str(i[0])))
                button = QRadioButton()
                self.savedReports.setIndexWidget(self.savedReports.model().index(n,2),button)
            for row in range(self.savedReports.rowCount()):
                for col in range(self.savedReports.columnCount()):
                    if col==3:
                        self.savedReports.item(row,col).setFlags(Qt.ItemFlag.ItemIsEditable)
                        
            createButton = QPushButton("عرض",self.windowSaved)
            createButton.clicked.connect(lambda:self.creating("Other"))
            createButton.setStyleSheet("background-color:green")
            createButton.setGeometry(150,460,150,30)

            extractAllButton = QPushButton("تصدير جميع التقارير",self.windowSaved)
            extractAllButton.clicked.connect(self.exportAllSummaryReports)
            extractAllButton.setStyleSheet("background-color:green")
            extractAllButton.setGeometry(320,460,160,30)

            self.windowSaved.show()
        def deleteReport(self,row,fRom="Original"):
            if fRom=="OutSide":
                d = QMessageBox(parent=self,text=f"تأكيد حذف تقرير {self.TableSummary.item(row,8).text()}")
            else:
                d = QMessageBox(parent=self,text=f"تأكيد حذف تقرير {self.savedReports.item(row,3).text()}")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle(title)
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                if fRom=="OutSide":
                    cr.execute(f"DELETE FROM reports WHERE id = '{self.TableSummary.item(row,8).text()}'")
                else:
                    cr.execute(f"DELETE FROM reports WHERE id = '{self.savedReports.item(row,0).text()}'")

                con.commit()
                if fRom=="OutSide":
                    self.TableSummary.hideRow(row)
                else:
                    self.savedReports.hideRow(row)
        def createReportFun(self):

            self.windowCreate = Choices()
            uic.loadUi("Create.ui",self.windowCreate)
            self.windowCreate.setFixedSize(359,370)
            self.windowCreate.setWindowTitle(title)
            self.windowCreate.setWindowIcon(QIcon("icon.ico"))
            self.windowCreate.CreateButton.setStyleSheet("background-color: cyan")
            self.windowCreate.CreateButton.clicked.connect(lambda:self.creating("Local"))
            self.windowCreate.CreateButton.setObjectName("ss")

            self.windowCreate.manyPic.addItem("بدون")
            self.windowCreate.manyPic.addItem("1")
            self.windowCreate.manyPic.addItem("2")
            self.windowCreate.manyPic.addItem("3")
            self.windowCreate.manyPic.addItem("4")

            self.windowCreate.NamePrograme.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Goals.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Description.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.WhenDate.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Benefits.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.BenefitsCount.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Creator.setCheckState(Qt.CheckState.Checked)
            

            self.windowCreate.show()
        def controlPanel(self):
            self.windowControl = Choices()
            self.windowControl.setFixedSize(300,570)
            self.windowControl.setWindowTitle(title)
            self.windowControl.setWindowIcon(QIcon("icon.ico"))
            
            self.picPathMinLogo = ""
            self.picBinaryMinLogo = ""
            
            layout = QVBoxLayout()
            self.windowControl.setLayout(layout)

            self.lineone = QLineEdit(self.windowControl)
            Label1 = QLabel("الترويسة الأولى",self.windowControl)
            Label1.setFont(QFont("Normal",15))
            Label1.move(198,10)
            self.lineone.setGeometry(7,40,290,30)
            
            self.linetwo = QLineEdit(self.windowControl)
            Label2 = QLabel("الترويسة الثانية",self.windowControl)
            Label2.setFont(QFont("Normal",15))
            Label2.move(198,70)
            self.linetwo.setGeometry(7,100,290,30)

            self.linethree = QLineEdit(self.windowControl)
            Label3 = QLabel("الترويسة الثالثة",self.windowControl)
            Label3.setFont(QFont("Normal",15))
            Label3.move(198,130)
            self.linethree.setGeometry(7,160,290,30)

            self.linefour = QLineEdit(self.windowControl)
            Label4 = QLabel("الترويسة الرابعة",self.windowControl)
            Label4.setFont(QFont("Normal",15))
            Label4.move(198,190)
            self.linefour.setGeometry(7,220,290,30)

            Label5 = QLabel("شعار وزارة التعليم",self.windowControl)
            Label5.setGeometry(90,250,120,30)

            self.FrameMin = QFrame(self.windowControl)
            self.FrameMin.setStyleSheet("background-color:rgb(178,178,178)")
            self.FrameMin.setGeometry(30,290,250,200)

            self.layoutFrameLogo = QVBoxLayout()
            self.FrameMin.setLayout(self.layoutFrameLogo)

            cr.execute("SELECT icon FROM start")

            with open("image12322.png","wb") as binary:
                binary.write(cr.fetchone()[0])
                
            img = Image.open("image12322.png")
            img = img.resize((250,170),Image.LANCZOS)
            img.save('image12322.png',quality=100)
            picLabel = QLabel(self.FrameMin)
            pix = QPixmap("image12322.png")
            picLabel.setPixmap(pix)

            self.layoutFrameLogo.addWidget(picLabel)

            specialButtonS = QPushButton(self.windowControl)
            specialButtonS.setStyleSheet(f"Qproperty-icon:url(cam.png);qproperty-iconSize:15px 15px;background-color:rgb(253, 253, 253)")
            specialButtonS.clicked.connect(self.addPicLogo)
            specialButtonS.setGeometry(130,490,30,30)


            self.SaveButton = QPushButton("حفظ",self.windowControl)
            self.SaveButton.setGeometry(55,530,200,30)


            cr.execute("SELECT * from start")
            values = cr.fetchall()[0]




            self.SaveButton.clicked.connect(self.Save)

            self.lineone.setText(values[0])
            self.linetwo.setText(values[1])
            self.linethree.setText(values[2])
            self.linefour.setText(values[3])
            self.windowControl.show()
        def addPicLogo(self):
            responce = QFileDialog.getOpenFileName(self.windowControl,"Select a File",desktopPath,filter="Image File (*.png *.jpg)")
            if len(responce[0])!=0:
                image = Image.open(responce[0])
                self.picPathMinLogo = responce[0]

                # finalImage = image.resize((350,200))
                finalImage = image.resize((250,170))

                finalImage.save("imagess22ss.png",quality=100)

                with open("imagess22ss.png","rb") as temp_binary:
                    binaryCode12 = temp_binary.read()

                self.picBinaryMinLogo = binaryCode12

                picLabel = QLabel(self.FrameMin)
                pix = QPixmap("imagess22ss.png")
                picLabel.setPixmap(pix)
                
                for i in reversed(range(self.layoutFrameLogo.count())): 
                    self.layoutFrameLogo.itemAt(i).widget().setParent(None)

                self.layoutFrameLogo.addWidget(picLabel)
                os.remove("imagess22ss.png")
        def Save(self):
            if self.picBinaryMinLogo !="":
                cr.execute(f"UPDATE start set line1='{self.lineone.text()}' ,line2='{self.linetwo.text()}' , line3='{self.linethree.text()}' , line4 = '{self.linefour.text()}',icon=?",([self.picBinaryMinLogo]))
            else:
                cr.execute(f"UPDATE start set line1='{self.lineone.text()}' ,line2='{self.linetwo.text()}' , line3='{self.linethree.text()}' , line4 = '{self.linefour.text()}'")
            con.commit()
            d = QMessageBox(parent=self,text="تم التعديل بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            ret = d.exec()
            self.windowControl.destroy()
        def creating(self,fromW):
            try:
                self.destroy()
                self.close()
                self.windowCreating.destroy()
                self.windowCreating.close()
            except:
                pass
            self.ablePrograme= False
            self.ableGoals= False
            self.ableDescription= False
            self.ableCreator= False
            self.ableDate= False
            self.ableBenefits= False
            self.ableCount= False
            self.countPic = 0
            self.secretLittleThing = ""
            if fromW=="Local":
                self.windowCreate.destroy()
            global yForImpo
            global xForImpo
            yForImpo = 50
            xForImpo = 0
            self.pictersPaths = ["","","",""]
            self.windowCreating = Choices()
            self.windowCreating.setMinimumSize(900,1150)
            self.setMinimumSize(930,500)
            self.setWindowTitle(title)
            self.windowCreating.setWindowTitle(title)
            self.windowCreating.setWindowIcon(QIcon("icon.ico"))
            # baseLayout = QVBoxLayout()
            # self.windowCreating.setLayout(baseLayout)

            self.hiderFrame = QFrame(self.windowCreating) #-> old way
            # self.hiderFrame = QFrame() # -> new way
            self.hiderFrame.setStyleSheet("background-color: white")
            # self.hiderFrame.setFixedSize(900,140)
            #-> old way

            self.hidderFramePic = QFrame(self.hiderFrame)
            self.hidderFramePic.setStyleSheet(f"background-color:#EBEAE9;")
            self.hidderFramePic.setGeometry(40,5,250,130)

            self.hidderlayoutPic = QVBoxLayout()
            self.hidderFramePic.setLayout(self.hidderlayoutPic)

            specialButton = QPushButton(self.hiderFrame)
            specialButton.setIcon(QIcon("cam.png"))
            specialButton.clicked.connect(lambda:self.putImage(f"lol"))
            specialButton.move(290,65)

            DeleteButtonHidder = QPushButton(self.hiderFrame)
            DeleteButtonHidder.setStyleSheet(f"Qproperty-icon:url(trashicon.png);qproperty-iconSize:15px 15px;background-color:rgb(253, 253, 253)")
            DeleteButtonHidder.clicked.connect(lambda:self.deleteImagesTemp(f"lol"))
            DeleteButtonHidder.move(290,90)

            



            Frame_text = QFrame(self.hiderFrame)
            Frame_text.setStyleSheet("background-color: white")
            # Frame_text.setFixedSize(250,140)
            Frame_text.setGeometry(655,0,250,140)

            text_layout = QVBoxLayout(Frame_text)

            Frame_text.setLayout(text_layout)
            cr.execute("SELECT line1 FROM start")
            Label1 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label1)
            
            cr.execute("SELECT line2 FROM start")
            Label2 = QLabel("   "+cr.fetchone()[0])
            text_layout.addWidget(Label2)

            cr.execute("SELECT line3 FROM start")
            Label3 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label3)

            cr.execute("SELECT line4 FROM start")
            Label4 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label4)

            logoLabel = QLabel(self.hiderFrame)
            logoLabel.move(350,10)
            cr.execute("SELECT icon FROM start")
            
            with open("logo.png","wb") as logoImpo:
                logoImpo.write(cr.fetchone()[0])
            img = Image.open("logo.png")
            img = img.resize((220,125),Image.LANCZOS)
            img.save("logo.png",quality=100)
            pix = QPixmap("logo.png")
            logoLabel.setPixmap(pix)
            # baseLayout.addWidget(self.hiderFrame) -> new way
            self.hiderFrame.setGeometry(0,3,900,140)
            
            self.cFrame = QFrame(self.windowCreating) ##-> oldway
            # self.cFrame = QFrame()
            self.cFrame.setStyleSheet("background-color: white")
            # self.cFrame.setMinimumSize(900,1100)
            emptyFields = QPushButton(self.cFrame,clicked=self.emptyFieldsFun)
            emptyFields.setToolTip('افراغ الحقول')
            emptyFields.setStyleSheet("font-size:12px;qproperty-icon:url('clearFields.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            # emptyFields.clicked.connect(lambda:self.putImage(f"lol"))
            emptyFields.move(860,5)

            savedReports = QPushButton(self.cFrame)
            savedReports.setToolTip('التقارير المحفوظة')
            savedReports.setStyleSheet("font-size:12px;qproperty-icon:url('savedReports.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            savedReports.clicked.connect(self.savedReportsFun)
            savedReports.move(820,5)

            CreateAnewReport = QPushButton(self.cFrame)
            CreateAnewReport.setToolTip('انشاء تقرير جديد')
            CreateAnewReport.setStyleSheet("font-size:12px;qproperty-icon:url('addNewreport.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            CreateAnewReport.clicked.connect(self.createReportFun)
            CreateAnewReport.move(820,40)

            summary = QPushButton(self.cFrame)
            summary.setToolTip('ملخص التقارير')
            summary.setStyleSheet("font-size:12px;qproperty-icon:url('summary.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            summary.clicked.connect(self.summaryReports)
            summary.move(860,40)


            printButton = QPushButton(self.cFrame)
            printButton.setToolTip('طباعة')
            printButton.setStyleSheet("font-size:12px;qproperty-icon:url('printer.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            printButton.clicked.connect(self.printDoc)
            printButton.move(860,75)

            controlButton = QPushButton(self.cFrame)
            controlButton.setToolTip('لوحة التحكم')
            controlButton.setStyleSheet("font-size:12px;qproperty-icon:url('control.png');qproperty-iconSize:25px 25px;background-color:rgb(148, 148, 148);")
            controlButton.clicked.connect(self.controlPanel)
            controlButton.move(820,75)


            labelGood = QLabel("توثيق برنامج",self.cFrame)
            labelGood.setStyleSheet("font-size:20px")
            labelGood.move(410,20)



            buttonSavePdf = QPushButton(self.cFrame,clicked=lambda x,y="Pdf":self.writeWord(y))
            buttonSavePdf.setStyleSheet("qproperty-icon:url(pdfIcon.png);qproperty-iconSize:40px 40px;")
            buttonSavePdf.move(10,10)

            buttonSaveWord = QPushButton(self.cFrame,clicked=self.writeWord)
            buttonSaveWord.setStyleSheet("qproperty-icon:url(word.png);qproperty-iconSize:40px 40px;")
            buttonSaveWord.move(65,10)

            buttonSaveInSidePrograme = QPushButton("حفظ بإسم",self.cFrame,clicked=self.SavePrograme)
            buttonSaveInSidePrograme.setStyleSheet("font-size:13px;background-color:cyan")
            buttonSaveInSidePrograme.move(10,60)



            if fromW=="Other":
                update = QPushButton("حفظ",self.cFrame,clicked=self.updateAReport)
                update.setStyleSheet("font-size:13px;background-color:cyan")
                update.move(10,90)

            self.programeNameShow = ""
            if fromW=="Local":
                self.programeNameShow = "ss12323"
            if fromW!="Local":
                numberOfPictures = -1
                for i in range(self.savedReports.rowCount()):
                    if self.savedReports.cellWidget(i,2).isChecked():
                        self.programeNameShow = self.savedReports.item(i,0).text()
                if self.programeNameShow !="":
                    cr.execute(f"SELECT * FROM reports WHERE id = '{self.programeNameShow}'")
                    listImportant = cr.fetchall()[0]
                    if listImportant[2]!="":
                        self.ablePrograme = True
                        self.createNamePrograme()
                        if listImportant[2] == " ":
                            self.programeNameE.setText(str(listImportant[2]).strip())
                        else:
                            self.programeNameE.setText(str(listImportant[2]))
                    if listImportant[3]!="":
                        self.ableGoals = True
                        self.createGoals()
                        self.programeGoalsE.setText(listImportant[3])
                        if listImportant[3] == " ":
                            self.programeGoalsE.setText(str(listImportant[3]).strip())
                        else:
                            self.programeGoalsE.setText(str(listImportant[3]))

                    if listImportant[4]!="":
                        self.ableDescription = True
                        self.createDescription()
                        self.programeDescriptionE.setText(listImportant[4])
                        if listImportant[4] == " ":
                            self.programeDescriptionE.setText(str(listImportant[4]).strip())
                        else:
                            self.programeDescriptionE.setText(str(listImportant[4]))


                    if listImportant[5]!="":
                        self.ableCreator = True
                        self.executer()
                        self.programeCreatorE.setText(listImportant[5])
                        if listImportant[5] == " ":
                            self.programeCreatorE.setText(str(listImportant[5]).strip())
                        else:
                            self.programeCreatorE.setText(str(listImportant[5]))
                    if listImportant[6]!="":
                        self.ableDate = True
                        self.executeDate()
                        self.programeWhenDateE.setText(listImportant[6])
                        if listImportant[6] == " ":
                            self.programeWhenDateE.setText(str(listImportant[6]).strip())
                        else:
                            self.programeWhenDateE.setText(str(listImportant[6]))

                    if listImportant[7]!="":
                        self.ableBenefits = True
                        self.Benefits()
                        self.programeBenefitsE.setText(listImportant[7])
                        if listImportant[7] == " ":
                            self.programeBenefitsE.setText(str(listImportant[7]).strip())
                        else:
                            self.programeBenefitsE.setText(str(listImportant[7]))
                    if listImportant[8]!="":
                        self.ableCount = True
                        self.CountBenefits()
                        self.CountBenefitsE.setText(listImportant[8])
                        if listImportant[8] == " ":
                            self.CountBenefitsE.setText(str(listImportant[8]).strip())
                        else:
                            self.CountBenefitsE.setText(str(listImportant[8]))
                    if listImportant[9]!="":
                        numberOfPictures+=1
                    if listImportant[10]!="":
                        numberOfPictures+=1  
                    if listImportant[11]!="":
                        numberOfPictures+=1  
                    if listImportant[12]!="":
                        numberOfPictures+=1
                    if numberOfPictures > -1:
                        self.CreatePic(numberOfPictures+1)
                        self.countPic = numberOfPictures+1
                    if listImportant[9]!="" and listImportant[9]!=" ":
                        with open(f"pic1.png","wb") as image:
                            image.write(listImportant[9])
                        self.pictersPaths[0] = "pic1.png"
                        self.putImage("Other",0)
                    if listImportant[10]!="" and listImportant[10]!=" ":
                        with open(f"pic2.png","wb") as image:
                            image.write(listImportant[10])
                        self.pictersPaths[1] = "pic2.png"
                        self.putImage("Other",1)
                    if listImportant[11]!="" and listImportant[12]!=" ":
                        with open(f"pic3.png","wb") as image:
                            image.write(listImportant[11])
                        self.pictersPaths[2] = "pic3.png"
                        self.putImage("Other",2)
                    if listImportant[12]!="" and listImportant[12]!=" ":
                        with open(f"pic4.png","wb") as image:
                            image.write(listImportant[12])
                        self.pictersPaths[3] = "pic4.png"
                        self.putImage("Other",3)
                    if listImportant[13]!="":
                        self.secretLittleThing =listImportant[13]
                        self.picLogoBinary = listImportant[13]
                        with open(f"secretThing.png","wb") as image:
                            image.write(self.secretLittleThing)
                        image = Image.open("secretThing.png")
                        finalImage = image.resize((350,200))

                        finalImage.save("secretThing.png",quality=100)
                        self.secretLittleThing = "secretThing.png"
                        self.putImage("Other",1000)
                    self.addAdmins()
                    if listImportant[14]!="":
                        self.label1Maye.setText(listImportant[14])
                    if listImportant[15]!="":
                        self.label2Maye.setText(listImportant[15])

                    if listImportant[16]!="":
                        self.MangerName.setText(listImportant[16])
                    if listImportant[17]!="":
                        self.consultName.setText(listImportant[17])

            else:
                if self.windowCreate.NamePrograme.isChecked():
                    self.ablePrograme = True
                    self.createNamePrograme()
                if self.windowCreate.Goals.isChecked():
                    self.ableGoals = True
                    self.createGoals()
                if self.windowCreate.Description.isChecked():
                    self.ableDescription = True
                    self.createDescription()
                if self.windowCreate.Creator.isChecked():
                    self.ableCreator = True
                    self.executer()
                if self.windowCreate.WhenDate.isChecked():
                    self.ableDate = True
                    self.executeDate()
                if self.windowCreate.Benefits.isChecked():
                    self.ableBenefits = True
                    self.Benefits()
                if self.windowCreate.BenefitsCount.isChecked():
                    self.ableCount = True
                    self.CountBenefits()
                if self.windowCreate.manyPic.currentText() != "بدون":
                    self.countPic = self.windowCreate.manyPic.currentText()
                    self.CreatePic(self.windowCreate.manyPic.currentText())
                self.addAdmins()
            self.cFrame.setGeometry(0,145,900,1150)
            # baseLayout.addWidget(self.cFrame)
            self.scroll = QScrollArea()             # Scroll Area which contains the widgets, set as the centralWidget
            #Scroll Area Properties
            self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            self.scroll.setWidgetResizable(True)
            self.scroll.setWidget(self.windowCreating)
            self.setCentralWidget(self.scroll)
            self.show()
            if self.programeNameShow =="":
                self.close()
                self.destroy()
                self.windowCreating.close()
                self.windowCreating.destroy()
        def summaryReports(self):
            self.windowsummary = Choices()
            self.windowsummary.resize(720,500)
            self.windowsummary.setWindowTitle(title)
            self.windowsummary.setWindowIcon(QIcon("icon.ico"))

            self.TableSummary = QTableWidget(self.windowsummary)
            self.TableSummary.setColumnCount(9)
            self.TableSummary.setHorizontalHeaderLabels(["","عدد المستفيدين","المستفيدون","تاريخ التنفيذ","المنفذ","الوصف","الأهداف","اسم البرنامج",""])
            self.TableSummary.setColumnHidden(8,True)
            self.windowsummary.resizeEvent = self.resizeSummary
            self.TableSummary.setColumnWidth(0,30)


            cr.execute("SELECT id,name,Goals,description,executer,executeDate,benefits,countBenefits From reports")
            result = cr.fetchall()
            for r,i in enumerate(result):
                self.TableSummary.insertRow(self.TableSummary.rowCount())
                i = list(i)
                i.insert(8,"")
                for col,c in enumerate(reversed(i)):
                    if col == 0:
                        button = QPushButton()
                        button.setStyleSheet(f"Qproperty-icon:url(trashIcon.png);qproperty-iconSize:30px 30px;background-color:rgb(253, 253, 253)")
                        button.clicked.connect(lambda x,row=r:self.deleteReport(row,"OutSide"))
                        self.TableSummary.setIndexWidget(self.TableSummary.model().index(r,0),button)
                    else:
                        item = QTableWidgetItem(str(c))
                        self.TableSummary.setItem(r,col,item)
            delegate = LineEditDelegate()
            self.TableSummary.setItemDelegate(delegate)
            self.TableSummary.setGeometry(0,0,720,470)
            self.TableSummary.cellDoubleClicked.connect(self.zoomSumarry)
            
            self.pdfExport = QPushButton("Pdf تصدير",self.windowsummary,clicked=lambda:self.exportSummaryScreen("Pdf"))
            self.pdfExport.setStyleSheet("background-color:red;font-size:20px")
            self.pdfExport.setGeometry(0,470,360,30)

            self.WordExport = QPushButton("Word تصدير",self.windowsummary,clicked=self.exportSummaryScreen)
            self.WordExport.setStyleSheet("background-color:blue;font-size:20px")
            self.WordExport.setGeometry(self.pdfExport.width(),470,360,30)

            self.priwidth = 720
            self.prihei = 500
            self.windowsummary.show()
        def zoomSumarry(self,row,col):
            if col !=0 and col!=8 :
                self.windowsummaryZoom = Choices()
                self.windowsummaryZoom.setFixedSize(400,300)
                self.windowsummaryZoom.setWindowTitle(title)
                self.windowsummaryZoom.setWindowIcon(QIcon("icon.ico"))
                
                summaryZoomText = QTextEdit(self.windowsummaryZoom)
                summaryZoomText.setGeometry(0,0,400,300)
                summaryZoomText.setAlignment(Qt.AlignmentFlag.AlignLeft)
                summaryZoomText.setFont(QFont("Arial",15))
                
                summaryZoomText.setText(self.TableSummary.item(row,col).text())

                self.windowsummaryZoom.show()
        def resizeSummary(self,ev):
            self.widthChanged = self.windowsummary.width() - self.priwidth
            self.heightChanged = self.windowsummary.height() - self.prihei - 10
            each = self.widthChanged-5

            
            self.TableSummary.resize(self.TableSummary.width()+self.widthChanged,self.TableSummary.height()+self.heightChanged)

            self.pdfExport.resize(self.pdfExport.width()+(round(self.widthChanged/2)),self.pdfExport.height())
            self.pdfExport.move(0,self.TableSummary.height()+3)
            
            self.WordExport.resize(self.WordExport.width()+(round(self.widthChanged/2)),self.WordExport.height())
            self.WordExport.move(self.pdfExport.width(),self.TableSummary.height()+3)


            incread = ((self.TableSummary.width() - 29) // 7) - 4
            if incread >= 87:
                for i in range(self.TableSummary.columnCount()):
                    if i!=0 and i!=8:
                        self.TableSummary.setColumnWidth(i,incread)


            self.priwidth =self.windowsummary.width()
            self.prihei = self.windowsummary.height()- 10


            self.TableSummary.show()
        def updateAReport(self):
            namePrograme = ""
            Goals = ""
            description = ""
            executer = ""
            executeDate = ""
            benefits = ""
            countBenefits = ""
            pic1 = ""
            pic2 = ""
            pic3 = ""
            pic4 = ""
            if self.ablePrograme:
                namePrograme = self.programeNameE.toPlainText() if len(self.programeNameE.toPlainText()) > 0 else " "
            if self.ableGoals:
                Goals = self.programeGoalsE.toPlainText() if len(self.programeGoalsE.toPlainText()) > 0 else " "
            if self.ableDescription:
                description = self.programeDescriptionE.toPlainText() if len(self.programeDescriptionE.toPlainText()) > 0 else " "

            if self.ableCreator:
                executer = self.programeCreatorE.toPlainText() if len(self.programeCreatorE.toPlainText()) > 0 else " "

            if self.ableDate:
                executeDate = self.programeWhenDateE.toPlainText() if len(self.programeWhenDateE.toPlainText()) > 0 else " "

            if self.ableBenefits:
                benefits = self.programeBenefitsE.toPlainText() if len(self.programeBenefitsE.toPlainText()) > 0 else " "

            if self.ableCount:
                countBenefits = self.CountBenefitsE.toPlainText() if len(self.CountBenefitsE.toPlainText()) > 0 else " "

            if self.countPic != 0:
                if self.pictersPaths[0]!="":
                    with open(self.pictersPaths[0],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic1 = binaryCode
                else:
                    pic1 = " "
                if self.pictersPaths[1]!="":
                    with open(self.pictersPaths[1],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic2 = binaryCode
                else:
                    pic2 = " "
                if self.pictersPaths[2]!="":
                    with open(self.pictersPaths[2],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic3 = binaryCode
                else:
                    pic3 = " "
                if self.pictersPaths[3]!="":
                    with open(self.pictersPaths[3],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic4 = binaryCode
                else:
                    pic4 = " "
            picLogo = ""
            if self.hidderlayoutPic.count() > 0:
                picLogo = self.picLogoBinary
            reportName = self.programeNameShow
            label1Maybe = self.label1Maye.text()
            label2Maybe = self.label2Maye.text()
            manger = str(self.MangerName.text())
            co_manger = str(self.consultName.text())
            cr.execute(f"""UPDATE reports set name=?,Goals=?,description=?,executer=?,executeDate=?,benefits=?,countBenefits=?,pic1=?,pic2=?,pic3=?,pic4=?,picLogo=?,label1Maybe=?,label2Maybe=?,manger=?,co_manger=? WHERE id = ?""",(namePrograme,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger,reportName))
            # it doesn't work just this way i have to use it to work
            d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            ret = d.exec()
            con.commit()
        def deleteImagesTemp(self,neNum):
            if neNum=="lol":
                for i in reversed(range(self.hidderlayoutPic.count())): 
                    self.hidderlayoutPic.itemAt(i).widget().setParent(None)
                    self.secretLittleThing = ""
            elif neNum==0:
                for i in reversed(range(self.layouts[0].count())): 
                    self.layouts[0].itemAt(i).widget().setParent(None)
                    self.pictersPaths[0] = ""
                    
            elif neNum==1:
                for i in reversed(range(self.layouts[1].count())): 
                    self.layouts[1].itemAt(i).widget().setParent(None)
                    self.pictersPaths[1] = ""

            elif neNum==2:
                for i in reversed(range(self.layouts[2].count())): 
                    self.layouts[2].itemAt(i).widget().setParent(None)
                    self.pictersPaths[2] = ""

            elif neNum==3:
                for i in reversed(range(self.layouts[3].count())): 
                    self.layouts[3].itemAt(i).widget().setParent(None)
                    self.pictersPaths[3] = ""
        def addAdmins(self):
            self.label1Maye = QLineEdit(self.cFrame)
            self.label1Maye.setGeometry(40,900,180,25)

            self.consultName = QLineEdit(self.cFrame)
            self.consultName.setGeometry(40,930,180,25)

            self.label2Maye = QLineEdit(self.cFrame)
            self.label2Maye.setGeometry(680,900,180,25)

            self.MangerName = QLineEdit(self.cFrame)
            self.MangerName.setGeometry(680,930,180,25)
        def SavePrograme(self):
            self.saveProgrameWindow = Choices()
            self.saveProgrameWindow.setFixedSize(200,200)
            self.saveProgrameWindow.setWindowTitle(title)
            self.saveProgrameWindow.setWindowIcon(QIcon("icon.ico"))
            self.saveProgrameWindow.setStyleSheet("background-color:white")
            
            Label = QLabel("اسم الملف",self.saveProgrameWindow)
            Label.move(20,20)

            self.NameEntryProgrameFile = QLineEdit(self.saveProgrameWindow)
            self.NameEntryProgrameFile.setFont(QFont("Arial",15))
            self.NameEntryProgrameFile.move(20,40)

            SaveButton = QPushButton("حفظ",self.saveProgrameWindow,clicked=self.saveReport)
            SaveButton.setGeometry(0,0,150,40)
            SaveButton.setStyleSheet("background-color:green")
            SaveButton.setFont(QFont("Arial",15))
            SaveButton.move(20,130)
            self.saveProgrameWindow.show()
        def CreatePic(self,Count):
            x = 100
            y = 440
            self.picters = []
            self.buttons = []
            self.layouts = []
            for i in range(int(Count)):
                if i > 1:
                    if i==2:
                        x = 100
                        y+=230
                    else:
                        x+=360
                self.picters.append(QFrame(self.cFrame))
                self.picters[i].setGeometry(0,0,350,200)
                self.picters[i].setStyleSheet(f"background-color:#EBEAE9;")
                self.picters[i].move(x,y)
                templayout = QVBoxLayout()
                self.picters[i].setLayout(templayout)
                self.layouts.append(templayout)
                button = QPushButton(self.cFrame)
                tempvar = self.picters[i].geometry()
                buttonx = tempvar.x()
                buttony = tempvar.y()
                self.buttons.append(button)
                self.buttons[i].move(buttonx+tempvar.width()//2 - 10,buttony+tempvar.height())
                self.buttons[i].setIcon(QIcon("cam.png"))
                self.buttons[i].setObjectName(f"{i}")
                self.buttons[i].clicked.connect(lambda ch,i=i:self.putImage(f"{i}"))

                DeleteButtonHidderInside = QPushButton(self.cFrame)
                DeleteButtonHidderInside.setStyleSheet(f"Qproperty-icon:url(trashicon.png);qproperty-iconSize:15px 16px;background-color:rgb(253, 253, 253)")
                DeleteButtonHidderInside.clicked.connect(lambda ch,x=i:self.deleteImagesTemp(x))
                DeleteButtonHidderInside.move(buttonx+tempvar.width()//2 - 35,buttony+tempvar.height())
                if i<=1:
                    x+=360
        def putImage(self,ob,number=-1):
            # self.picters[int(ob)].dele
            try:
                os.remove("image.png")
            except:
                pass
            if ob=="lol":
                responce = QFileDialog.getOpenFileName(self.windowCreating,"Select a File",desktopPath,filter="Image File (*.png *.jpg)")
                if len(responce[0])!=0:
                    image = Image.open(responce[0])
                    self.secretLittleThing = responce[0]

                    # finalImage = image.resize((350,200))
                    finalImage = image.resize((240,110))

                    finalImage.save("image.png",quality=100)

                    with open("image.png","rb") as temp_binary:
                        binaryCode12 = temp_binary.read()

                    self.picLogoBinary = binaryCode12

                    picLabel = QLabel(self.hidderFramePic)
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    
                    for i in reversed(range(self.hidderlayoutPic.count())): 
                        self.hidderlayoutPic.itemAt(i).widget().setParent(None)

                    self.hidderlayoutPic.addWidget(picLabel)
                    os.remove("image.png")
            elif ob=="Other":
                if number==0:
                    cr.execute(f"SELECT pic1 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==1:
                    cr.execute(f"SELECT pic2 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==2:
                    cr.execute(f"SELECT pic3 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==3:
                    cr.execute(f"SELECT pic4 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==1000:
                    cr.execute(f"SELECT picLogo FROM reports WHERE id = '{self.programeNameShow}'")
                # image = Image.open(responce[0])
                with open(f"pic1.png","wb") as image:
                    image.write(cr.fetchone()[0])
                image = Image.open("pic1.png")
                if number==1000:
                    finalImage = image.resize((240,110))
                else:
                    # finalImage = image.resize((240,180))
                    finalImage = image.resize((350,180))

                finalImage.save("image.png",quality=100)
                if number==0:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==1:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==2:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==3:
                    picLabel = QLabel(self.hidderFramePic)
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==1000:
                    picLabel = QLabel(self.hidderFramePic)
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.hidderlayoutPic.count())): 
                        self.hidderlayoutPic.itemAt(i).widget().setParent(None)
                    self.hidderlayoutPic.addWidget(picLabel)
                # os.remove("image.png")
                # try:
                #     os.remove("pic1.png")
                # except:
                #     pass
            else:
                responce = QFileDialog.getOpenFileName(self.windowCreating,"Select a File",desktopPath,filter="Image File (*.png *.jpg)")
                if len(responce[0])!=0:
                    try:
                        os.remove("image.png")
                    except:
                        pass
                    self.pictersPaths[int(ob)]=(responce[0])
                    image = Image.open(responce[0])
                    # finalImage = image.resize((240,180))
                    finalImage = image.resize((350,180))

                    finalImage.save("image.png",quality=100)
                    picLabel = QLabel(self.picters[int(ob)])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    # self.picters[int(ob)].
                    # Playout.
                    for i in reversed(range(self.layouts[int(ob)].count())): 
                        self.layouts[int(ob)].itemAt(i).widget().setParent(None)
                    #change the wat to put the image use the one in the teachers programe 
                    self.layouts[int(ob)].addWidget(picLabel)
                    os.remove("image.png")
        def createNamePrograme(self):
            global yForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("اسم البرنامج")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeNameE = QTextEdit(self.cFrame)
            self.programeNameE.setGeometry(0,0,565,35)
            self.programeNameE.setAlignment(Qt.AlignmentFlag.AlignLeft)

            self.programeNameE.setFont(QFont("Arial",14))
            self.programeNameE.move(135,yForImpo)


            yForImpo +=35
        def resizedWindow(self,ev):
            #self.cFrame self.hiderFrame
            newWidth = (self.width() - self.cFrame.width())//2
            self.cFrame.move(newWidth,self.cFrame.y())
            self.hiderFrame.move(newWidth,self.hiderFrame.y())

        def createGoals(self):
            global yForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("\n\n  الاهداف ")
            programeName.setGeometry(0,0,100,110)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeGoalsE = QTextEdit(self.cFrame)
            self.programeGoalsE.setGeometry(0,0,565,110)
            self.programeGoalsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeGoalsE.setFont(QFont("Arial",15))
            self.programeGoalsE.move(135,yForImpo) # 280
            yForImpo+=110
        def createDescription(self):
            global yForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("\n\n  الوصف ")
            programeName.setGeometry(0,0,100,110)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeDescriptionE = QTextEdit(self.cFrame)
            self.programeDescriptionE.setGeometry(0,0,565,110)
            self.programeDescriptionE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeDescriptionE.setFont(QFont("Arial",15))
            self.programeDescriptionE.move(135,yForImpo)
            yForImpo+=110
        def executer(self):
            global yForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("المنفذ")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeCreatorE = QTextEdit(self.cFrame)
            self.programeCreatorE.setGeometry(10,10,565,35)
            self.programeCreatorE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeCreatorE.setFont(QFont("Arial",15))
            self.programeCreatorE.move(135,yForImpo)
            yForImpo+=35
        def executeDate(self):
            global yForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("تاريخ التنفيذ")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            
            
            self.programeWhenDateE = QTextEdit(self.cFrame)
            self.programeWhenDateE.setGeometry(10,10,565,35)
            self.programeWhenDateE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeWhenDateE.setFont(QFont("Arial",14))
            self.programeWhenDateE.move(135,yForImpo)
            yForImpo+=35
        def Benefits(self):
            global yForImpo
            global xForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("المستفيدون")
            programeName.setGeometry(0,0,100,30)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",13))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeBenefitsE = QTextEdit(self.cFrame)
            self.programeBenefitsE.setGeometry(10,10,565,30)
            self.programeBenefitsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeBenefitsE.setFont(QFont("Arial",13))
            self.programeBenefitsE.move(135,yForImpo)
            yForImpo += 30
        def emptyFieldsFun(self):
            d = QMessageBox(parent=self.windowCreating,text="تأكيد افراغ جميع الحقول")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle(title)
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                try:
                    self.programeNameE.setText("")
                except:
                    pass
                try:
                    self.programeNameE.setText("")
                except:
                    pass
                try:
                    self.programeGoalsE.setText("")
                except:
                    pass
                try:
                    self.programeDescriptionE.setText("")
                except:
                    pass
                try:
                    self.programeCreatorE.setText("")
                except:
                    pass
                try:
                    self.programeWhenDateE.setText("")
                except:
                    pass
                try:
                    self.programeBenefitsE.setText("")
                except:
                    pass
                try:
                    self.CountBenefitsE.setText("")
                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[0].count())): 
                        self.layouts[0].itemAt(i).widget().setParent(None)
                        self.pictersPaths[0] = ""
                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[1].count())): 
                        self.layouts[1].itemAt(i).widget().setParent(None)
                        self.pictersPaths[1] = ""

                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[2].count())): 
                        self.layouts[2].itemAt(i).widget().setParent(None)
                        self.pictersPaths[2] = ""

                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[3].count())): 
                        self.layouts[3].itemAt(i).widget().setParent(None)
                        self.pictersPaths[3] = ""
                except:
                    pass

                try:
                    self.label1Maye.setText("")
                except:
                    pass
                try:
                    self.label2Maye.setText("")
                except:
                    pass

                try:
                    self.consultName.setText("")
                except:
                    pass

                try:
                    self.MangerName.setText("")
                except:
                    pass
        def CountBenefits(self):
            global yForImpo
            global xForImpo
            programeName = QTextEdit(self.cFrame)
            programeName.setText("عدد المستفيدين")
            programeName.setGeometry(0,0,100,30)
            programeName.setStyleSheet("background-color: #2ABCB5")

            programeName.setFont(QFont("Arial",13))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            





            self.CountBenefitsE = QTextEdit(self.cFrame)
            self.CountBenefitsE .setGeometry(10,10,565,30)
            self.CountBenefitsE .setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.CountBenefitsE .setFont(QFont("Arial",13))

            self.CountBenefitsE.move(135,yForImpo)
        def saveReport(self):
            namePrograme = ""
            Goals = ""
            description = ""
            executer = ""
            executeDate = ""
            benefits = ""
            countBenefits = ""
            pic1 = ""
            pic2 = ""
            pic3 = ""
            pic4 = ""
            if self.ablePrograme:
                namePrograme = self.programeNameE.toPlainText() if len(self.programeNameE.toPlainText()) > 0 else " "
            if self.ableGoals:
                Goals = self.programeGoalsE.toPlainText() if len(self.programeGoalsE.toPlainText()) > 0 else " "
            if self.ableDescription:
                description = self.programeDescriptionE.toPlainText() if len(self.programeDescriptionE.toPlainText()) > 0 else " "

            if self.ableCreator:
                executer = self.programeCreatorE.toPlainText() if len(self.programeCreatorE.toPlainText()) > 0 else " "

            if self.ableDate:
                executeDate = self.programeWhenDateE.toPlainText() if len(self.programeWhenDateE.toPlainText()) > 0 else " "

            if self.ableBenefits:
                benefits = self.programeBenefitsE.toPlainText() if len(self.programeBenefitsE.toPlainText()) > 0 else " "

            if self.ableCount:
                countBenefits = self.CountBenefitsE.toPlainText() if len(self.CountBenefitsE.toPlainText()) > 0 else " "
            if self.countPic != 0:
                if self.pictersPaths[0]!="":
                    with open(self.pictersPaths[0],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic1 = binaryCode
                else:
                    pic1 = " "
                if self.pictersPaths[1]!="":
                    with open(self.pictersPaths[1],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic2 = binaryCode
                else:
                    pic2 = " "
                if self.pictersPaths[2]!="":
                    with open(self.pictersPaths[2],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic3 = binaryCode
                else:
                    pic3 = " "
                if self.pictersPaths[3]!="":
                    with open(self.pictersPaths[3],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic4 = binaryCode
                else:
                    pic4 = " "
            picLogo = ""
            if self.hidderlayoutPic.count() > 0:
                picLogo = self.picLogoBinary
            reportName = str(self.NameEntryProgrameFile.text())
            label1Maybe = self.label1Maye.text()
            label2Maybe = self.label2Maye.text()
            manger = str(self.MangerName.text())
            co_manger = str(self.consultName.text())
            cr.execute(f"""Insert INTO reports (reportName,name,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(reportName,namePrograme,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger))
            d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            ret = d.exec()
            self.saveProgrameWindow.destroy()
            con.commit()
        def writeWord(self,fromWhere="Word"):
            FileNameSave = QFileDialog.getSaveFileName(self.windowCreating,"Select File",desktopPath)
            if len(FileNameSave[0])>0:
                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                doc = docx.Document()
                sections = doc.sections
                for section in sections:
                    section.top_margin = docx.shared.Cm(0.7)
                    section.bottom_margin = docx.shared.Cm(0.7)
                    section.left_margin = docx.shared.Cm(0.7)
                    section.right_margin = docx.shared.Cm(0.7)
                sec_pr = doc.sections[0]._sectPr # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'triple') # a single line
                    border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '10')
                    border_el.set(qn('w:color'), 'black')
                    pg_borders.append(border_el) # register single border to border el
                sec_pr.append(pg_borders) # apply border changes to section

                headers_table = doc.add_table(rows=1, cols=2)
                for row in headers_table.rows:
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        tc.left = None
                        tc.top = None
                        tc.right = None
                        tc.bottom = None
                        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                        cell.paragraphs[0].size = docx.shared.Pt(15)

                hdr_Cells = headers_table.rows[0].cells

                cr.execute("SELECT line1 FROM start")
                hdr_Cells[1].text = cr.fetchone()[0]
                cr.execute("SELECT line2 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]+"\t"
                cr.execute("SELECT line3 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]
                cr.execute("SELECT line4 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]

                widths = (docx.shared.Inches(5.8),docx.shared.Inches(3))
                for row in headers_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                heights = (docx.shared.Inches(1.1),docx.shared.Inches(1.1))
                for idx,row in enumerate(headers_table.rows):
                    row.height = heights[idx]

                paragraph12322 =hdr_Cells[1].paragraphs[0]
                run = paragraph12322.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)

                cells = headers_table.rows[0].cells[0].paragraphs[0]
                runCells = cells.add_run()
                if self.secretLittleThing !="":
                    runCells.add_picture(self.secretLittleThing,width=docx.shared.Inches(2.1),height=docx.shared.Inches(1))
                if self.hidderlayoutPic.count() <= 0:
                    runCells.add_text("\t\t\t\t\t")
                else:
                    runCells.add_text("\t")
                xsaw = runCells.add_picture("logo.png",width=docx.shared.Inches(2.5),height=docx.shared.Inches(1))

                for row in headers_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                headers_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


                """
                if self.windowCreate.BenefitsCount.isChecked():
                    self.CountBenefits()
                """

                GoodPrograme = doc.add_paragraph("\t\t\t\t\t\t\tتوثيق برنامج")
                GoodPrograme.runs[0].font.size = docx.shared.Pt(20)
                GoodPrograme.paragraph_format.space_after = docx.shared.Pt(0.1)
                GoodPrograme.paragraph_format.space_before = docx.shared.Pt(1)

                if self.ablePrograme:
                    text = self.programeNameE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeNameProgrameTable = doc.add_table(rows=1,cols=2)

                    programeNameProgrameTable.style = "Table Grid"
                    hdr_Cells = programeNameProgrameTable.rows[0].cells
                    hdr_Cells[1].text = "اسم البرنامج"
                    hdr_Cells[0].text = "".join(final_text)
                    programeNameProgrameTable.autofit = False

                    cell_xml_element = programeNameProgrameTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))

                    for row in programeNameProgrameTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    
                    for row in programeNameProgrameTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(17)
                    heights = (docx.shared.Inches(.35), docx.shared.Inches(.35))
                    for idx,row in enumerate(programeNameProgrameTable.rows):
                            row.height = heights[idx]
                if self.ableGoals:
                    text = self.programeGoalsE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeGolasTable = doc.add_table(rows=1,cols=2)

                    programeGolasTable.style = "Table Grid"
                    hdr_Cells = programeGolasTable.rows[0].cells
                    hdr_Cells[1].text = "\t\tالأهداف"
                    hdr_Cells[0].text = "".join(final_text)
                    programeGolasTable.autofit = False

                    cell_xml_element = programeGolasTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))
                    for row in programeGolasTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    for idx,row in enumerate(programeGolasTable.rows):
                        row.height = heights[idx]
                    

                    for row in programeGolasTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)


                if self.ableDescription:
                    text = self.programeDescriptionE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeDescriptionTable = doc.add_table(rows=1,cols=2)

                    programeDescriptionTable.style = "Table Grid"
                    hdr_Cells = programeDescriptionTable.rows[0].cells
                    hdr_Cells[1].text = "\t\tالوصف"
                    hdr_Cells[0].text = "".join(final_text)
                    programeDescriptionTable.autofit = False

                    cell_xml_element = programeDescriptionTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(0.9))
                    heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))

                    for row in programeDescriptionTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width

                    for idx,row in enumerate(programeDescriptionTable.rows):
                            row.height = heights[idx]

                    for row in programeDescriptionTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)

                if self.ableCreator:
                    text = self.programeCreatorE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeCreatorTable = doc.add_table(rows=1,cols=2)

                    programeCreatorTable.style = "Table Grid"
                    hdr_Cells = programeCreatorTable.rows[0].cells

                    hdr_Cells[0].text = "".join(final_text)
                    hdr_Cells[1].text = "المنفذ"
                    programeCreatorTable.autofit = False


                    cell_xml_element = programeCreatorTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeCreatorTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width

                    for row in programeCreatorTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)

            
                if self.ableDate:
                    text = self.programeWhenDateE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeWhenDateTable = doc.add_table(rows=1,cols=2)

                    programeWhenDateTable.style = "Table Grid"
                    hdr_Cells = programeWhenDateTable.rows[0].cells
                    hdr_Cells[1].text = "تاريخ التنفيذ"
                    hdr_Cells[0].text = "".join(final_text)
                    programeWhenDateTable.autofit = False

                    cell_xml_element = programeWhenDateTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeWhenDateTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    
                    for row in programeWhenDateTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(17)
                if self.ableBenefits:
                    text = self.programeBenefitsE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)
                            
                    programeBenefitsTable = doc.add_table(rows=1,cols=2)
                    programeBenefitsTable.style = 'Table Grid' #single lines in all cells
                    hdr_Cells = programeBenefitsTable.rows[0].cells
                    hdr_Cells[1].text = "المستفيدون"
                    hdr_Cells[0].text = "".join(final_text)
                    programeBenefitsTable.autofit = False


                    cell_xml_element = programeBenefitsTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeBenefitsTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width            
                    for row in programeBenefitsTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(15)
                if self.ableCount:
                    text = self.programeNameE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeCountBenefitsTable = doc.add_table(rows=1,cols=2)
                    programeCountBenefitsTable.style = 'Table Grid' 
                    hdr_Cells = programeCountBenefitsTable.rows[0].cells
                    hdr_Cells[1].text = "عدد المستفيدين"
                    hdr_Cells[0].text = "".join(final_text)
                    programeCountBenefitsTable.autofit = False

                    cell_xml_element = programeCountBenefitsTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeCountBenefitsTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width            
                    for row in programeCountBenefitsTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(14)

                if self.countPic != 0:
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.space_before = docx.shared.Pt(2)
                    run = paragraph.add_run()
                    for i in range(len(self.pictersPaths)):
                        if self.pictersPaths[i] !="":
                                try:
                                    os.remove("imageWithBoarder.png")
                                except:
                                    pass
                                imgB = Image.open(self.pictersPaths[i])

                                border_color_rgb = (128, 128, 128, 255)

                                resize = imgB.resize((500,500),Image.LANCZOS)

                                bordered_image = ImageOps.expand(resize, border=8, fill=border_color_rgb)
                                
                                bordered_image.save('imageWithBoarder.png')
                                if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                                    run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(1.8))
                                else:
                                    run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(2.2))

                                if i !=1:
                                    run.add_text("   ")
                                if i==1:
                                    run.add_text("\n")
                        paragraph.paragraph_format.space_after = docx.shared.Pt(0)

                if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                    addmins_table = doc.add_table(rows=1, cols=2)
                    for row in addmins_table.rows:
                        for cell in row.cells:
                            tc = cell._element.tcPr
                            tc.left = None
                            tc.top = None
                            tc.right = None
                            tc.bottom = None
                            cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                            cell.paragraphs[0].size = docx.shared.Pt(8)

                    addmins_Cells = addmins_table.rows[0].cells

                    addmins_Cells[0].text = self.label1Maye.text()+"\n"+f"{self.consultName.text()}"

                    addmins_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                    addmins_Cells[1].text = self.label2Maye.text()+"\n"+f"{self.MangerName.text()}"

                    addmins_table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                    paragraph12 =addmins_Cells[0].paragraphs[0]
                    run = paragraph12.runs
                    font = run[0].font
                    font.size= docx.shared.Pt(15)

                    paragraph13 =addmins_Cells[1].paragraphs[0]
                    run = paragraph13.runs
                    font = run[0].font
                    font.size= docx.shared.Pt(15)
                    heights = (docx.shared.Pt(16),docx.shared.Pt(16))
                    for idx,row in enumerate(addmins_table.rows):
                        row.height = heights[idx]
                subFilesD = [f for f in os.listdir(folderFinle) if f.endswith(".docx")]
                name2 = nameFile+".docx"
                if name2 in subFilesD:
                    i = 1
                    while name2 in subFilesD:
                        name2 = f"({i}) {name2}"
                        i+=1
                


                
                doc.save(f"{folderFinle}/{name2}")
                s = False
                if fromWhere =="Pdf":
                    subFilesD = [f for f in os.listdir(folderFinle) if f.endswith(".pdf")]
                    name3 = nameFile+".pdf"
                    if name3 in subFilesD:
                        i = 1
                        while name3 in subFilesD:
                            name3 = f"({i}) {name3}"
                            i+=1

                with suppress_output():
                    convert(f"{folderFinle}/{name2}",f"{folderFinle}/{name3}")
                os.remove(f"{folderFinle}/{name2}")
                d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                ret = d.exec()
            try:
                os.remove("pic1")
                os.remove("pic2")
                os.remove("pic4")
                os.remove("pic4")
                os.remove("secretThing.png")
            except:
                pass
        def printDoc(self,pdf):
            try:
                os.remove("printFile.pdf")
                os.remove("printFile.docx")
            except:
                pass
            doc = docx.Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = docx.shared.Cm(0.7)
                section.bottom_margin = docx.shared.Cm(0.7)
                section.left_margin = docx.shared.Cm(0.7)
                section.right_margin = docx.shared.Cm(0.7)
            sec_pr = doc.sections[0]._sectPr # get the section properties el
            # create new borders el
            pg_borders = OxmlElement('w:pgBorders')
            # specifies how the relative positioning of the borders should be calculated
            pg_borders.set(qn('w:offsetFrom'), 'page')
            for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'triple') # a single line
                border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
                border_el.set(qn('w:space'), '10')
                border_el.set(qn('w:color'), 'black')
                pg_borders.append(border_el) # register single border to border el
            sec_pr.append(pg_borders) # apply border changes to section

            headers_table = doc.add_table(rows=1, cols=2)
            for row in headers_table.rows:
                for cell in row.cells:
                    tc = cell._element.tcPr
                    tc.left = None
                    tc.top = None
                    tc.right = None
                    tc.bottom = None
                    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell.paragraphs[0].size = docx.shared.Pt(15)
            
            hdr_Cells = headers_table.rows[0].cells

            cr.execute("SELECT line1 FROM start")
            hdr_Cells[1].text = cr.fetchone()[0]
            cr.execute("SELECT line2 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]+"\t"
            cr.execute("SELECT line3 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]
            cr.execute("SELECT line4 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]

            widths = (docx.shared.Inches(5.8),docx.shared.Inches(3))
            for row in headers_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

            heights = (docx.shared.Inches(1.1),docx.shared.Inches(1.1))
            for idx,row in enumerate(headers_table.rows):
                row.height = heights[idx]

            paragraph12322 =hdr_Cells[1].paragraphs[0]
            run = paragraph12322.runs
            font = run[0].font
            font.size= docx.shared.Pt(15)

            cells = headers_table.rows[0].cells[0].paragraphs[0]
            runCells = cells.add_run()
            if self.secretLittleThing !="":
                runCells.add_picture(self.secretLittleThing,width=docx.shared.Inches(2.1),height=docx.shared.Inches(1))
            if self.hidderlayoutPic.count() <= 0:
                runCells.add_text("\t\t\t\t\t")
            else:
                runCells.add_text("\t")
            xsaw = runCells.add_picture("logo.png",width=docx.shared.Inches(2.5),height=docx.shared.Inches(1))

            for row in headers_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            headers_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


            """
            if self.windowCreate.BenefitsCount.isChecked():
                self.CountBenefits()
            """

            GoodPrograme = doc.add_paragraph("\t\t\t\t\t\t\tتوثيق برنامج")
            GoodPrograme.runs[0].font.size = docx.shared.Pt(20)
            GoodPrograme.paragraph_format.space_after = docx.shared.Pt(0.1)
            GoodPrograme.paragraph_format.space_before = docx.shared.Pt(1)



            if self.ablePrograme:
                text = self.programeNameE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeNameProgrameTable = doc.add_table(rows=1,cols=2)

                programeNameProgrameTable.style = "Table Grid"
                hdr_Cells = programeNameProgrameTable.rows[0].cells
                hdr_Cells[1].text = "اسم البرنامج"
                hdr_Cells[0].text = ''.join(final_text)
                programeNameProgrameTable.autofit = False

                cell_xml_element = programeNameProgrameTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))

                for row in programeNameProgrameTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeNameProgrameTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                heights = (docx.shared.Inches(.35), docx.shared.Inches(.35))
                for idx,row in enumerate(programeNameProgrameTable.rows):
                        row.height = heights[idx]
            if self.ableGoals:
                text = self.programeGoalsE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeGolasTable = doc.add_table(rows=1,cols=2)

                programeGolasTable.style = "Table Grid"
                hdr_Cells = programeGolasTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالأهداف"
                hdr_Cells[0].text = ''.join(final_text)
                programeGolasTable.autofit = False

                cell_xml_element = programeGolasTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))
                for row in programeGolasTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                for idx,row in enumerate(programeGolasTable.rows):
                    row.height = heights[idx]
                

                for row in programeGolasTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)

            if self.ableDescription:
                text = self.programeDescriptionE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeDescriptionTable = doc.add_table(rows=1,cols=2)

                programeDescriptionTable.style = "Table Grid"
                hdr_Cells = programeDescriptionTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالوصف"
                hdr_Cells[0].text = ''.join(final_text)
                programeDescriptionTable.autofit = False

                cell_xml_element = programeDescriptionTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(0.9))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))

                for row in programeDescriptionTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for idx,row in enumerate(programeDescriptionTable.rows):
                        row.height = heights[idx]

                for row in programeDescriptionTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)

            if self.ableCreator:
                text = self.programeCreatorE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeCreatorTable = doc.add_table(rows=1,cols=2)

                programeCreatorTable.style = "Table Grid"
                hdr_Cells = programeCreatorTable.rows[0].cells

                hdr_Cells[0].text = ''.join(final_text)
                hdr_Cells[1].text = "المنفذ"
                programeCreatorTable.autofit = False


                cell_xml_element = programeCreatorTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCreatorTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for row in programeCreatorTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)

        
            if self.ableDate:
                text = self.programeWhenDateE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeWhenDateTable = doc.add_table(rows=1,cols=2)

                programeWhenDateTable.style = "Table Grid"
                hdr_Cells = programeWhenDateTable.rows[0].cells
                hdr_Cells[1].text = "تاريخ التنفيذ"
                hdr_Cells[0].text = ''.join(final_text)
                programeWhenDateTable.autofit = False

                cell_xml_element = programeWhenDateTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeWhenDateTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeWhenDateTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
            if self.ableBenefits:
                text = self.programeBenefitsE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeBenefitsTable = doc.add_table(rows=1,cols=2)
                programeBenefitsTable.style = 'Table Grid' #single lines in all cells
                hdr_Cells = programeBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "المستفيدون"
                hdr_Cells[0].text = ''.join(final_text)
                programeBenefitsTable.autofit = False


                cell_xml_element = programeBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(15)
            if self.ableCount:
                text = self.CountBenefitsE.toPlainText()
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeCountBenefitsTable = doc.add_table(rows=1,cols=2)
                programeCountBenefitsTable.style = 'Table Grid' 
                hdr_Cells = programeCountBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "عدد المستفيدين"
                hdr_Cells[0].text = ''.join(final_text)
                programeCountBenefitsTable.autofit = False

                cell_xml_element = programeCountBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCountBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeCountBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(14)

            if self.countPic != 0:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = docx.shared.Pt(2)
                run = paragraph.add_run()
                for i in range(len(self.pictersPaths)):
                    if self.pictersPaths[i] !="":
                            try:
                                os.remove("imageWithBoarder.png")
                            except:
                                pass
                            imgB = Image.open(self.pictersPaths[i])

                            border_color_rgb = (128, 128, 128, 255)

                            resize = imgB.resize((500,500),Image.LANCZOS)

                            bordered_image = ImageOps.expand(resize, border=8, fill=border_color_rgb)
                            
                            bordered_image.save('imageWithBoarder.png')
                            if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                                run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(1.8))
                            else:
                                run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(2.2))

                            if i !=1:
                                run.add_text("   ")
                            if i==1:
                                run.add_text("\n")
                    paragraph.paragraph_format.space_after = docx.shared.Pt(0)

            if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                addmins_table = doc.add_table(rows=1, cols=2)
                for row in addmins_table.rows:
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        tc.left = None
                        tc.top = None
                        tc.right = None
                        tc.bottom = None
                        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                        cell.paragraphs[0].size = docx.shared.Pt(8)
                addmins_Cells = addmins_table.rows[0].cells

                addmins_Cells[0].text = self.label1Maye.text()+"\n"+f"{self.consultName.text()}"

                addmins_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                addmins_Cells[1].text = self.label2Maye.text()+"\n"+f"{self.MangerName.text()}"

                addmins_table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                paragraph12 =addmins_Cells[0].paragraphs[0]
                run = paragraph12.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)

                paragraph13 =addmins_Cells[1].paragraphs[0]
                run = paragraph13.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)
                heights = (docx.shared.Pt(16),docx.shared.Pt(16))
                for idx,row in enumerate(addmins_table.rows):
                    row.height = heights[idx]

            doc.save("printFile.docx")

            with suppress_output():
                convert("printFile.docx","printFile.pdf")
            os.remove("printFile.docx")
            
            try:
                os.remove("pic1")
                os.remove("pic2")
                os.remove("pic4")
                os.remove("pic4")
                os.remove("secretThing.png")
            except:
                pass
            # os.startfile("printFile.pdf")
            webbrowser.open("printFile.pdf", new=2)

            time.sleep(5)
            pyautogui.hotkey("ctrl","p")
        def exportAllSummaryReports(self):
            FileNameSave = QFileDialog.getSaveFileName(self.windowSaved,"Select File",desktopPath)
            self.pdfFilesPaths = []
            if len(FileNameSave[0])>0:
                files = os.listdir("tempPdf")
                for file in files:
                    file_path = os.path.join("tempPdf",file)
                    if os.path.isfile(file_path):
                        os.remove(file_path)

                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                cr.execute("SELECT id FROM reports")
                ides = cr.fetchall()
                self.eachValue = 100//len(ides)
                self.progressBarWindow = Choices()
                self.progressBarWindow.setFixedSize(250,30)
                self.progressBar = QProgressBar(self.progressBarWindow)
                self.progressBar.setGeometry(0,0,290,30)
                self.progressBarWindow.show()
                for i in ides:
                    for j in i:
                        self.completeExportAllSummaryReports(j)

                merger = PdfMerger()
                for pdf in self.pdfFilesPaths:
                    merger.append(open(pdf, 'rb'))
                with open(f"{FileNameSave[0]}.pdf", "wb") as fout:
                    merger.write(fout)
                self.progressBarWindow.destroy()
                d = QMessageBox(parent=None,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                ret = d.exec()
        def completeExportAllSummaryReports(self,idFun):
            try:
                os.remove("pic1.png")
                os.remove("pic2.png")
                os.remove("pic4.png")
                os.remove("pic4.png")
                os.remove("secretThing.png")
            except:
                pass

            doc = docx.Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = docx.shared.Cm(0.7)
                section.bottom_margin = docx.shared.Cm(0.7)
                section.left_margin = docx.shared.Cm(0.7)
                section.right_margin = docx.shared.Cm(0.7)
            sec_pr = doc.sections[0]._sectPr # get the section properties el
            # create new borders el
            pg_borders = OxmlElement('w:pgBorders')
            # specifies how the relative positioning of the borders should be calculated
            pg_borders.set(qn('w:offsetFrom'), 'page')
            for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'triple') # a single line
                border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
                border_el.set(qn('w:space'), '10')
                border_el.set(qn('w:color'), 'black')
                pg_borders.append(border_el) # register single border to border el
            sec_pr.append(pg_borders) # apply border changes to section

            headers_table = doc.add_table(rows=1, cols=2)
            for row in headers_table.rows:
                for cell in row.cells:
                    tc = cell._element.tcPr
                    tc.left = None
                    tc.top = None
                    tc.right = None
                    tc.bottom = None
                    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell.paragraphs[0].size = docx.shared.Pt(15)
            hdr_Cells = headers_table.rows[0].cells

            cr.execute("SELECT line1 FROM start")
            hdr_Cells[1].text = cr.fetchone()[0]
            cr.execute("SELECT line2 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]+"\t"
            cr.execute("SELECT line3 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]
            cr.execute("SELECT line4 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]

            widths = (docx.shared.Inches(5.8),docx.shared.Inches(3))
            for row in headers_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

            heights = (docx.shared.Inches(1.1),docx.shared.Inches(1.1))
            for idx,row in enumerate(headers_table.rows):
                row.height = heights[idx]

            paragraph12322 =hdr_Cells[1].paragraphs[0]
            run = paragraph12322.runs
            font = run[0].font
            font.size= docx.shared.Pt(15)

            cells = headers_table.rows[0].cells[0].paragraphs[0]
            runCells = cells.add_run()
            cr.execute(f"SELECT picLogo FROM reports WHERE id={idFun}")
            picLogo = cr.fetchone()[0]
            if picLogo !="":
                with open("secretThing.png","wb") as secretThing:
                    secretThing.write(picLogo)
                runCells.add_picture("secretThing.png",width=docx.shared.Inches(2.1),height=docx.shared.Inches(1))
                runCells.add_text("\t")
            else:
                runCells.add_text("\t\t\t\t\t")

            xsaw = runCells.add_picture("logo.png",width=docx.shared.Inches(2.5),height=docx.shared.Inches(1))

            for row in headers_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            headers_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


            GoodPrograme = doc.add_paragraph("\t\t\t\t\t\t\tتوثيق برنامج")
            GoodPrograme.runs[0].font.size = docx.shared.Pt(20)
            GoodPrograme.paragraph_format.space_after = docx.shared.Pt(0.1)
            GoodPrograme.paragraph_format.space_before = docx.shared.Pt(1)







            cr.execute(f"SELECT name FROM reports WHERE id={idFun}")  
            namePrograme = cr.fetchone()[0]
            if namePrograme!="":
                text = namePrograme
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeNameProgrameTable = doc.add_table(rows=1,cols=2)

                programeNameProgrameTable.style = "Table Grid"
                hdr_Cells = programeNameProgrameTable.rows[0].cells
                hdr_Cells[1].text = "اسم البرنامج"
                hdr_Cells[0].text = ''.join(final_text)
                programeNameProgrameTable.autofit = False

                cell_xml_element = programeNameProgrameTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))

                for row in programeNameProgrameTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeNameProgrameTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                heights = (docx.shared.Inches(.35), docx.shared.Inches(.35))
                for idx,row in enumerate(programeNameProgrameTable.rows):
                        row.height = heights[idx]
            
            cr.execute(f"SELECT Goals FROM reports WHERE id={idFun}")  
            Goals = cr.fetchone()[0]
            if Goals!="":
                text = Goals
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeGolasTable = doc.add_table(rows=1,cols=2)

                programeGolasTable.style = "Table Grid"
                hdr_Cells = programeGolasTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالأهداف"
                hdr_Cells[0].text = ''.join(final_text)
                programeGolasTable.autofit = False

                cell_xml_element = programeGolasTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))
                for row in programeGolasTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                for idx,row in enumerate(programeGolasTable.rows):
                    row.height = heights[idx]
                

                for row in programeGolasTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)
            
            cr.execute(f"SELECT description FROM reports WHERE id={idFun}")  
            description = cr.fetchone()[0]
            if description!="":
                text = description
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeDescriptionTable = doc.add_table(rows=1,cols=2)

                programeDescriptionTable.style = "Table Grid"
                hdr_Cells = programeDescriptionTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالوصف"
                hdr_Cells[0].text = ''.join(final_text)
                programeDescriptionTable.autofit = False

                cell_xml_element = programeDescriptionTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(0.9))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))

                for row in programeDescriptionTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for idx,row in enumerate(programeDescriptionTable.rows):
                        row.height = heights[idx]

                for row in programeDescriptionTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)
            
            cr.execute(f"SELECT executer FROM reports WHERE id={idFun}")  
            executer = cr.fetchone()[0]
            if executer!="":
                text = executer
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeCreatorTable = doc.add_table(rows=1,cols=2)

                programeCreatorTable.style = "Table Grid"
                hdr_Cells = programeCreatorTable.rows[0].cells

                hdr_Cells[0].text = ''.join(final_text)
                hdr_Cells[1].text = "المنفذ"
                programeCreatorTable.autofit = False


                cell_xml_element = programeCreatorTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCreatorTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for row in programeCreatorTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)

            cr.execute(f"SELECT executeDate FROM reports WHERE id={idFun}")
            executeDate = cr.fetchone()[0]
            if executeDate!="":
                text = executeDate
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeWhenDateTable = doc.add_table(rows=1,cols=2)

                programeWhenDateTable.style = "Table Grid"
                hdr_Cells = programeWhenDateTable.rows[0].cells
                hdr_Cells[1].text = "تاريخ التنفيذ"
                hdr_Cells[0].text = ''.join(final_text)
                programeWhenDateTable.autofit = False

                cell_xml_element = programeWhenDateTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeWhenDateTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeWhenDateTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
            
            cr.execute(f"SELECT benefits FROM reports WHERE id={idFun}")
            benefits = cr.fetchone()[0]

            if benefits!="":
                text = benefits
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)
                
                programeBenefitsTable = doc.add_table(rows=1,cols=2)
                programeBenefitsTable.style = 'Table Grid' #single lines in all cells
                hdr_Cells = programeBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "المستفيدون"
                hdr_Cells[0].text = ''.join(final_text)
                programeBenefitsTable.autofit = False


                cell_xml_element = programeBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(15)
            
            cr.execute(f"SELECT countBenefits FROM reports WHERE id={idFun}")
            countBenefits = cr.fetchone()[0]
            
            if countBenefits!="":
                text = countBenefits
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeCountBenefitsTable = doc.add_table(rows=1,cols=2)
                programeCountBenefitsTable.style = 'Table Grid' 
                hdr_Cells = programeCountBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "عدد المستفيدين"
                hdr_Cells[0].text = ''.join(final_text)
                programeCountBenefitsTable.autofit = False

                cell_xml_element = programeCountBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCountBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeCountBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(14)


            self.pictersPaths = []

            cr.execute(f"SELECT pic1 FROM reports WHERE id={idFun}")
            pic1B = cr.fetchone()[0]
            if pic1B!="":
                with open("pic1.png","wb") as pic1:
                    pic1.write(pic1B)
                self.pictersPaths.append("pic1.png")

            cr.execute(f"SELECT pic2 FROM reports WHERE id={idFun}")
            pic2B = cr.fetchone()[0]
            if pic2B!="":
                with open("pic2.png","wb") as pic2:
                    pic2.write(pic2B)
                self.pictersPaths.append("pic2.png")

            cr.execute(f"SELECT pic3 FROM reports WHERE id={idFun}")
            pic3B = cr.fetchone()[0]
            if pic3B!="":
                with open("pic3.png","wb") as pic3:
                    pic3.write(pic3B)
                self.pictersPaths.append("pic3.png")

            cr.execute(f"SELECT pic4 FROM reports WHERE id={idFun}")
            pic4B = cr.fetchone()[0]
            if pic4B!="":
                with open("pic4.png","wb") as pic4:
                    pic4.write(pic4B)
                self.pictersPaths.append("pic4.png")

            if len(self.pictersPaths) > 0:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = docx.shared.Pt(2)
                run = paragraph.add_run()
                for i in range(len(self.pictersPaths)):
                    try:
                        os.remove("imageWithBoarder.png")
                    except:
                        pass
                    imgB = Image.open(self.pictersPaths[i])

                    border_color_rgb = (128, 128, 128, 255)

                    resize = imgB.resize((500,500),Image.LANCZOS)

                    bordered_image = ImageOps.expand(resize, border=8, fill=border_color_rgb)
                    
                    bordered_image.save('imageWithBoarder.png')

                    cr.execute(f"SELECT label1Maybe FROM reports WHERE id={idFun}")
                    label1Maybe = cr.fetchone()[0]

                    cr.execute(f"SELECT label2Maybe FROM reports WHERE id={idFun}")
                    label2Maybe = cr.fetchone()[0]

                    if len(label1Maybe) > 0 or len(label2Maybe) > 0:
                        run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(1.8))
                    else:
                        run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(2.2))

                    if i !=1:
                        run.add_text("   ")
                    if i==1:
                        run.add_text("\n")
                    paragraph.paragraph_format.space_after = docx.shared.Pt(0)




            if len(label1Maybe) > 0 or len(label2Maybe) > 0:
                addmins_table = doc.add_table(rows=1, cols=2)
                for row in addmins_table.rows:
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        tc.left = None
                        tc.top = None
                        tc.right = None
                        tc.bottom = None
                        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                        cell.paragraphs[0].size = docx.shared.Pt(8)

                addmins_Cells = addmins_table.rows[0].cells


                cr.execute(f"SELECT manger FROM reports WHERE id={idFun}")
                manger = cr.fetchone()[0]


                addmins_Cells[0].text = label1Maybe+"\n"+f"{manger}"

                addmins_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                cr.execute(f"SELECT co_manger FROM reports WHERE id={idFun}")
                co_manger = cr.fetchone()[0]

                addmins_Cells[1].text = label2Maybe+"\n"+f"{co_manger}"

                addmins_table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                paragraph12 =addmins_Cells[0].paragraphs[0]
                run = paragraph12.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)

                paragraph13 =addmins_Cells[1].paragraphs[0]
                run = paragraph13.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)
                heights = (docx.shared.Pt(16),docx.shared.Pt(16))
                for idx,row in enumerate(addmins_table.rows):
                    row.height = heights[idx]

            name2 = str(idFun)+".docx"
            doc.save(f"tempPdf/{name2}")

            in_file = str(os.path.abspath(f"tempPdf/{name2}")).replace("c","C")
            out_file = str(os.path.abspath(f"tempPdf/{str(idFun)}.pdf")).replace("c","C")
            with suppress_output():
                convert(f"{in_file}",f"{out_file}")
            os.remove(f"tempPdf/{name2}")

            self.pdfFilesPaths.append(f"tempPdf/{str(idFun)}.pdf")
            self.progressBar.setValue(self.progressBar.value() + self.eachValue)
        def exportSummaryScreen(self,fromWhere="Word"):
            FileNameSave = QFileDialog.getSaveFileName(self.windowCreating,"Select File",desktopPath)
            if len(FileNameSave[0])>0:
                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                doc = docx.Document()
                sections = doc.sections
                for section in sections:
                    section.top_margin = docx.shared.Cm(0.7)
                    section.bottom_margin = docx.shared.Cm(0.7)
                    section.left_margin = docx.shared.Cm(0.7)
                    section.right_margin = docx.shared.Cm(0.7)

                SummryTable = doc.add_table(rows=1,cols=8)
                SummryTable.style = "Table Grid"
                hdr_Cells = SummryTable.rows[0].cells

                hdr_Cells[7].text = "م"
                hdr_Cells[6].text = "اسم البرنامج"
                hdr_Cells[5].text = "الأهداف"
                hdr_Cells[4].text = "الوصف"
                hdr_Cells[3].text = "المنفذ"
                hdr_Cells[2].text = "تاريخ التنفيذ"
                hdr_Cells[1].text = "المتسفيدون"
                hdr_Cells[0].text = "عدد المستفيدين"

                widths = (docx.shared.Inches(2), docx.shared.Inches(1.5),docx.shared.Inches(2),docx.shared.Inches(1.5),docx.shared.Inches(1.5),docx.shared.Inches(1.5),docx.shared.Inches(1.5),docx.shared.Inches(.4))

                cr.execute("SELECT name,Goals,description,executer,executeDate,benefits,countBenefits FROM reports")
                
                for numberTemp,i in enumerate(cr.fetchall()):
                    row_Cells = SummryTable.add_row().cells
                    row_Cells[7].text = str(numberTemp+1)
                    number = 6
                    for j in i:
                        text = j
                        listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                        final_text = []
                        for xs in text:
                            if xs in listNubmers:
                                final_text.append(convert_numbers.english_to_hindi(xs))
                            else:
                                final_text.append(xs)
                    
                        row_Cells[number].text = "".join(final_text)
                        number-=1
                for row in SummryTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in SummryTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)

                subFilesD = [f for f in os.listdir(folderFinle) if f.endswith(".docx")]
                name2 = nameFile+".docx"
                if name2 in subFilesD:
                    i = 1
                    while name2 in subFilesD:
                        name2 = f"({i}) {name2}"
                        i+=1
                
                doc.save(f"{folderFinle}/{name2}")
                s = False
                if fromWhere =="Pdf":
                    subFilesD = [f for f in os.listdir(folderFinle) if f.endswith(".pdf")]
                    name3 = nameFile+".pdf"
                    if name3 in subFilesD:
                        i = 1
                        while name3 in subFilesD:
                            name3 = f"({i}) {name3}"
                            i+=1


                    in_file = f"{folderFinle}/{name2}"
                    out_file = f"{folderFinle}/{name3}"

                    with suppress_output():
                        convert(in_file,out_file)

                    os.remove(f"{folderFinle}/{name2}")

                d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                ret = d.exec()
        def closeEvent(self, event):
            try:
                self.sender().objectName()
                event.accept()
            except:
                reply = QMessageBox()
                reply.setWindowTitle("تأكيد حفظ")
                reply.setText("هل تريد حفظ التقرير")

                reply.setStandardButtons(QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No|QMessageBox.StandardButton.Cancel)
                
                bottonOk = reply.button(QMessageBox.StandardButton.Yes)
                bottonOk.setText("نعم")
                
                bottonCancel = reply.button(QMessageBox.StandardButton.No)
                bottonCancel.setText("لا")
                
                bottonNo = reply.button(QMessageBox.StandardButton.Cancel)
                bottonNo.setText("تم الحفظ")

                x = reply.exec()

                if x == QMessageBox.StandardButton.No or x == QMessageBox.StandardButton.Cancel:
                    event.accept()
                elif x == QMessageBox.StandardButton.Yes:
                    event.ignore()
                    self.SavePrograme()        

    if __name__ == "__main__":
        app = QApplication(sys.argv)
        app.setStyleSheet(
            '''
            QLineEdit{
                font-size:15px
            }
            QLabel{
                font-size:15px
            }
            '''
        )
        window = Main()

        app.exec()
        
else:
    ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)

